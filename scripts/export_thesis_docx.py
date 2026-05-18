from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "paper" / "毕业论文初稿.md"
OUTPUT = ROOT / "paper" / "毕业论文初稿_排版稿.docx"


def set_run_font(run, size: float = 12, bold: bool = False, east_asia: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_font(paragraph, size: float = 12, bold: bool = False, east_asia: str = "宋体") -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, east_asia=east_asia)


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    return text


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)
    set_run_font(run, size=10.5)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.bold = True
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styles["Heading 3"].font.size = Pt(14)
    styles["Heading 3"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_heading(doc: Document, level: int, text: str) -> None:
    if text == "摘 要":
        text = "摘    要"
    if text == "致 谢":
        text = "致    谢"
    paragraph = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 14, bold=True, east_asia="黑体")


def add_normal_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, size=12, east_asia="宋体")


def add_center_caption(doc: Document, text: str, is_table: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.0 if not is_table else 1.5
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, size=10.5, east_asia="宋体")


def add_equation(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(" ".join(line.strip() for line in lines))
    set_run_font(run, size=12, east_asia="宋体")


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows: list[list[str]] = []
    for idx, line in enumerate(table_lines):
        cells = [clean_inline(cell.strip()) for cell in line.strip("|").split("|")]
        if idx == 1 and all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(len(table.columns)):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            run = paragraph.add_run(text)
            set_run_font(run, size=10.5, bold=(r_idx == 0), east_asia="宋体")


def add_image(doc: Document, image_path: Path) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run()
    if image_path.exists():
        run.add_picture(str(image_path), width=Cm(15.0))
    else:
        run.add_text(f"[图片缺失：{image_path}]")
        set_run_font(run, size=10.5, east_asia="宋体")


def export() -> None:
    doc = Document()
    configure_document(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_buffer: list[str] = []
    in_equation = False
    equation_buffer: list[str] = []
    i = 0

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.first_line_indent = None
                run = paragraph.add_run("\n".join(code_buffer))
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(9)
                code_buffer = []
                in_code = False
            else:
                in_code = True
                code_buffer = []
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if stripped == "$$":
            if in_equation:
                add_equation(doc, equation_buffer)
                equation_buffer = []
                in_equation = False
            else:
                in_equation = True
                equation_buffer = []
            i += 1
            continue
        if in_equation:
            equation_buffer.append(stripped)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            rows, next_i = parse_table(lines, i)
            add_table(doc, rows)
            i = next_i
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            image_rel = image_match.group(2)
            image_path = (SOURCE.parent / image_rel).resolve()
            add_image(doc, image_path)
            i += 1
            continue

        if stripped.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            run = paragraph.add_run(stripped[2:].strip())
            set_run_font(run, size=16, bold=True, east_asia="黑体")
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, 1, stripped[3:].strip())
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, 2, stripped[4:].strip())
            i += 1
            continue

        if re.match(r"^表\d+-\d+|^图\d+-\d+", stripped):
            add_center_caption(doc, stripped, is_table=stripped.startswith("表"))
            i += 1
            continue

        add_normal_paragraph(doc, stripped)
        i += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"已导出 Word 排版稿：{OUTPUT}")


if __name__ == "__main__":
    export()
